"""
Uploaded file -> plain text.

Any file the person has, in whatever form they keep it: a markdown journal, a
CV as PDF, a Word document, an exported chat log. Each one becomes text, and
the text is what the model reads.

Nothing here writes to disk. A life profile is the most personal thing this app
will ever touch, so it exists in memory for the length of a session and nowhere
else -- see `config.STORE_ON_DISK`.
"""
import io
import json
from pathlib import Path

from . import config


class UploadError(RuntimeError):
    """Raised when a file cannot be turned into text. Message is for the user."""


def _read_text(data: bytes, suffix: str) -> str:
    """Decode bytes that are already text, tolerating the usual encodings."""
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise UploadError("That file's text could not be decoded.")

    # A JSON export reads far better to the model pretty-printed than as one
    # enormous line, and it costs nothing to do here.
    if suffix == ".json":
        try:
            return json.dumps(json.loads(text), indent=2, ensure_ascii=False)
        except (json.JSONDecodeError, ValueError):
            return text
    return text


def _read_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise UploadError("PDF support needs `pypdf`. Run: pip install pypdf")

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception:
        raise UploadError(
            "That PDF could not be opened. It may be corrupt or password "
            "protected."
        )

    pages = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            pages.append(f"[page {number}]\n{text.strip()}")

    if not pages:
        # A scanned PDF is a picture of text, and pypdf finds nothing in it.
        # Saying so is more use than handing back an empty document.
        raise UploadError(
            "No text could be read from that PDF. If it is a scan or a photo, "
            "the words are an image rather than text — export a text-based PDF, "
            "or paste the content into a .md or .txt file instead."
        )
    return "\n\n".join(pages)


def _read_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise UploadError(
            "Word support needs `python-docx`. Run: pip install python-docx"
        )

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception:
        raise UploadError("That Word file could not be opened. It may be corrupt.")

    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # Tables often carry the actual facts in a CV, and reading only paragraphs
    # would silently drop them.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    if not parts:
        raise UploadError("That Word file has no readable text in it.")
    return "\n\n".join(parts)


def extract_text(data: bytes, filename: str) -> str:
    """One uploaded file -> its text. Raises UploadError with a readable reason."""
    if not data:
        raise UploadError(f"`{filename}` is empty.")
    if len(data) > config.MAX_UPLOAD_BYTES:
        raise UploadError(
            f"`{filename}` is {len(data) / 1_048_576:.1f} MB. The limit is "
            f"{config.MAX_UPLOAD_BYTES // 1_048_576} MB."
        )

    suffix = Path(filename).suffix.lower()

    if suffix in config.PDF_SUFFIXES:
        text = _read_pdf(data)
    elif suffix in config.DOCX_SUFFIXES:
        text = _read_docx(data)
    elif suffix in config.TEXT_SUFFIXES:
        text = _read_text(data, suffix)
    else:
        # Refuse by name rather than guessing. A .doc, .pages or .zip decoded as
        # latin-1 produces thousands of characters of rubbish that look like a
        # successful upload right up until the answers make no sense.
        raise UploadError(
            f"`{filename}` is not a file type this app can read. Supported: "
            f"{', '.join(sorted(config.SUPPORTED_SUFFIXES))}. "
            f"For anything else, copy the text into a .md or .txt file."
        )

    if not text.strip():
        raise UploadError(f"`{filename}` has no readable text in it.")
    return text.strip()


def build_profile(files) -> tuple:
    """`files` is a list of (filename, bytes) -> (profile_text, per_file_report).

    Every file is attempted. One unreadable file does not lose the others; it
    is reported next to them so the person can see what did and did not make it
    in, rather than wondering why an answer is missing something.
    """
    parts, report = [], []
    for filename, data in files:
        try:
            text = extract_text(data, filename)
        except UploadError as e:
            report.append({"filename": filename, "ok": False,
                           "chars": 0, "detail": str(e)})
            continue
        parts.append(f"===== FILE: {filename} =====\n{text}")
        report.append({"filename": filename, "ok": True,
                       "chars": len(text), "detail": None})

    profile = "\n\n".join(parts)

    truncated = False
    if len(profile) > config.MAX_PROFILE_CHARS:
        profile = profile[:config.MAX_PROFILE_CHARS]
        truncated = True

    return profile, {"files": report, "truncated": truncated,
                     "chars": len(profile)}
